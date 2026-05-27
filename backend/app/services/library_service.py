from __future__ import annotations

import base64
import csv
from datetime import datetime, timezone
from io import StringIO
import mimetypes
from pathlib import Path
import re
import uuid
from dataclasses import dataclass

from fastapi import HTTPException, UploadFile, status

from app.core.config import Settings
from app.core.database import ensure_storage, get_db
from app.schemas.chat import ChatAttachment
from app.schemas.library import LibraryItem


@dataclass
class AttachmentContext:
    attachment_id: str
    name: str
    type: str
    supported: bool
    excerpt: str | None = None
    note: str | None = None
    mime_type: str | None = None
    data_url: str | None = None


class LibraryService:
    def __init__(self, settings: Settings):
        self.settings = settings

    @staticmethod
    def _infer_type(name: str) -> str:
        extension = name.split(".")[-1].lower() if "." in name else ""
        if extension in {"png", "jpg", "jpeg", "gif", "webp"}:
            return "image"
        if extension == "pdf":
            return "pdf"
        if extension in {"ppt", "pptx"}:
            return "ppt"
        if extension in {"csv", "xls", "xlsx"}:
            return "csv"
        if extension in {"doc", "docx"}:
            return "doc"
        if extension in {"ts", "tsx", "js", "jsx", "py", "html", "css", "json"}:
            return "code"
        return "text"

    @staticmethod
    def _size_label(size: int) -> str:
        if size >= 1024 * 1024:
            return f"{size / (1024 * 1024):.1f} MB"
        return f"{max(1, round(size / 1024))} KB"

    @staticmethod
    def _row_to_item(row) -> LibraryItem:
        return LibraryItem(
            id=row["id"],
            owner_id=row["owner_id"],
            name=row["name"],
            type=row["type"],
            kind=row["kind"],
            source=row["source"],
            size_label=row["size_label"],
            created_at=row["created_at"],
        )

    def list_items(self, *, owner_id: str | None = None) -> list[LibraryItem]:
        query = """
            SELECT id, owner_id, name, type, kind, source, size_label, file_path, created_at
            FROM library_items
        """
        params: tuple = ()
        if owner_id:
            query += " WHERE owner_id = ?"
            params = (owner_id,)
        query += " ORDER BY datetime(created_at) DESC, name ASC"

        with get_db() as connection:
            rows = connection.execute(query, params).fetchall()
        return [self._row_to_item(row) for row in rows]

    def create_item(
        self,
        *,
        owner_id: str | None,
        name: str,
        content: bytes,
        source: str = "upload",
        kind: str = "file",
    ) -> LibraryItem:
        item_id = uuid.uuid4().hex
        created_at = datetime.now(timezone.utc).isoformat()
        _, upload_dir = ensure_storage(self.settings)
        extension = Path(name).suffix
        stored_name = f"{item_id}{extension}"
        file_path = upload_dir / stored_name
        file_path.write_bytes(content)

        item = LibraryItem(
            id=item_id,
            owner_id=owner_id,
            name=name,
            type=self._infer_type(name),
            kind=kind,
            source=source,
            size_label=self._size_label(len(content)),
            created_at=created_at,
        )

        with get_db() as connection:
            connection.execute(
                """
                INSERT INTO library_items
                (id, owner_id, name, type, kind, source, size_label, file_path, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item.id,
                    item.owner_id,
                    item.name,
                    item.type,
                    item.kind,
                    item.source,
                    item.size_label,
                    str(file_path),
                    item.created_at,
                ),
            )

        return item

    async def create_item_from_upload(
        self,
        *,
        owner_id: str | None,
        upload: UploadFile,
    ) -> LibraryItem:
        original_name = self._sanitize_file_name(upload.filename or "upload")
        extension = Path(original_name).suffix.lower()
        content_type = (upload.content_type or "").lower()

        if extension not in self.settings.allowed_upload_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="This file type is not allowed.",
            )
        generic_binary = content_type == "application/octet-stream"
        generic_text = content_type.startswith("text/")
        if (
            content_type
            and content_type not in self.settings.allowed_upload_content_types
            and not generic_binary
            and not generic_text
        ):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="This content type is not allowed.",
            )

        item_id = uuid.uuid4().hex
        created_at = datetime.now(timezone.utc).isoformat()
        _, upload_dir = ensure_storage(self.settings)
        stored_name = f"{item_id}{extension}"
        storage_path = upload_dir / stored_name
        total_bytes = 0

        with storage_path.open("wb") as output:
            while True:
                chunk = await upload.read(1024 * 1024)
                if not chunk:
                    break
                total_bytes += len(chunk)
                if total_bytes > self.settings.max_upload_size_bytes:
                    output.close()
                    storage_path.unlink(missing_ok=True)
                    raise HTTPException(
                        status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                        detail="Uploaded file exceeds the maximum allowed size.",
                    )
                output.write(chunk)

        await upload.close()

        if total_bytes == 0:
            storage_path.unlink(missing_ok=True)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Uploaded file is empty.",
            )

        item = LibraryItem(
            id=item_id,
            owner_id=owner_id,
            name=original_name,
            type=self._infer_type(original_name),
            kind="file",
            source="upload",
            size_label=self._size_label(total_bytes),
            created_at=created_at,
        )

        with get_db() as connection:
            connection.execute(
                """
                INSERT INTO library_items
                (id, owner_id, name, type, kind, source, size_label, file_path, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    item.id,
                    item.owner_id,
                    item.name,
                    item.type,
                    item.kind,
                    item.source,
                    item.size_label,
                    str(storage_path),
                    item.created_at,
                ),
            )

        return item

    def build_attachment_context(
        self,
        *,
        attachments: list[ChatAttachment],
        owner_id: str | None,
    ) -> list[AttachmentContext]:
        if not attachments:
            return []

        attachment_ids = [attachment.id for attachment in attachments if not attachment.id.startswith("tool-")]
        if not attachment_ids:
            return []

        placeholders = ",".join("?" for _ in attachment_ids)
        query = f"""
            SELECT id, owner_id, name, type, file_path
            FROM library_items
            WHERE id IN ({placeholders})
        """
        params: list[str] = attachment_ids[:]
        if owner_id:
            query += " AND owner_id = ?"
            params.append(owner_id)

        with get_db() as connection:
            rows = connection.execute(query, tuple(params)).fetchall()

        rows_by_id = {row["id"]: row for row in rows}
        contexts: list[AttachmentContext] = []

        for attachment in attachments:
            if attachment.id.startswith("tool-"):
                continue

            row = rows_by_id.get(attachment.id)
            if not row:
                contexts.append(
                    AttachmentContext(
                        attachment_id=attachment.id,
                        name=attachment.name,
                        type=attachment.type,
                        supported=False,
                        note="Attachment metadata was received, but the file could not be resolved from the library.",
                    )
                )
                continue

            file_path = row["file_path"]
            if not file_path:
                contexts.append(
                    AttachmentContext(
                        attachment_id=attachment.id,
                        name=row["name"],
                        type=row["type"],
                        supported=False,
                        note="The file is registered but no local payload is available.",
                    )
                )
                continue

            path = Path(file_path)
            if not path.exists():
                contexts.append(
                    AttachmentContext(
                        attachment_id=attachment.id,
                        name=row["name"],
                        type=row["type"],
                        supported=False,
                        note="The file record exists, but the stored payload is missing on disk.",
                    )
                )
                continue

            path = Path(file_path)

            if row["type"] == "image":
                mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
                contexts.append(
                    AttachmentContext(
                        attachment_id=attachment.id,
                        name=row["name"],
                        type=row["type"],
                        supported=True,
                        note="Image attachment is available for vision-capable models.",
                        mime_type=mime_type,
                        data_url=f"data:{mime_type};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}",
                    )
                )
                continue

            extracted_text, extract_note = self._extract_readable_text(path, row["type"])
            if not extracted_text:
                contexts.append(
                    AttachmentContext(
                        attachment_id=attachment.id,
                        name=row["name"],
                        type=row["type"],
                        supported=False,
                        note=extract_note
                        or (
                            "The file is attached, but its content could not be converted into readable text."
                        ),
                    )
                )
                continue

            contexts.append(
                AttachmentContext(
                    attachment_id=attachment.id,
                    name=row["name"],
                    type=row["type"],
                    supported=True,
                    excerpt=extracted_text[:6000].strip(),
                    note=extract_note,
                )
            )

        return contexts

    def _extract_readable_text(self, path: Path, item_type: str) -> tuple[str, str | None]:
        suffix = path.suffix.lower()

        if item_type in {"text", "code"} or suffix in {".txt", ".md", ".json", ".py", ".ts", ".tsx", ".js", ".jsx", ".html", ".css"}:
            return self._decode_text_payload(path.read_bytes()[:32000]), None

        if suffix == ".csv":
            decoded = self._decode_text_payload(path.read_bytes()[:64000])
            if not decoded:
                return "", "The CSV file could not be decoded as readable text."
            return self._summarize_csv(decoded), None

        if suffix in {".xlsx"}:
            return self._extract_xlsx(path)

        if suffix == ".xls":
            return "", "Legacy .xls files are not supported yet. Please upload CSV or XLSX."

        if suffix == ".pdf":
            return self._extract_pdf(path)

        if suffix == ".docx":
            return self._extract_docx(path)

        if suffix == ".doc":
            return "", "Legacy .doc files are not supported yet. Please upload DOCX, PDF, or TXT."

        return "", f"This file type is attached but not readable yet: {suffix or item_type}."

    @staticmethod
    def _summarize_csv(decoded: str) -> str:
        rows = list(csv.reader(StringIO(decoded)))
        if not rows:
            return ""
        preview = rows[:40]
        return "\n".join(", ".join(cell.strip() for cell in row[:20]) for row in preview)

    @staticmethod
    def _extract_pdf(path: Path) -> tuple[str, str | None]:
        try:
            from pypdf import PdfReader
        except Exception:  # noqa: BLE001
            return "", "PDF parsing dependency is not installed."

        try:
            reader = PdfReader(str(path))
            pages: list[str] = []
            for page in reader.pages[:20]:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(text.strip())
            return "\n\n".join(pages), None if pages else "No readable text was found in the PDF."
        except Exception as exc:  # noqa: BLE001
            return "", f"PDF could not be parsed: {exc}"

    @staticmethod
    def _extract_docx(path: Path) -> tuple[str, str | None]:
        try:
            from docx import Document
        except Exception:  # noqa: BLE001
            return "", "DOCX parsing dependency is not installed."

        try:
            document = Document(str(path))
            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            table_lines: list[str] = []
            for table in document.tables[:8]:
                for row in table.rows[:20]:
                    table_lines.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            text = "\n".join([*paragraphs, *table_lines])
            return text, None if text.strip() else "No readable text was found in the DOCX file."
        except Exception as exc:  # noqa: BLE001
            return "", f"DOCX could not be parsed: {exc}"

    @staticmethod
    def _extract_xlsx(path: Path) -> tuple[str, str | None]:
        try:
            from openpyxl import load_workbook
        except Exception:  # noqa: BLE001
            return "", "XLSX parsing dependency is not installed."

        try:
            workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets[:5]:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(min_row=1, max_row=40, max_col=20, values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(value.strip() for value in values):
                        lines.append(" | ".join(values))
            workbook.close()
            text = "\n".join(lines)
            return text, None if text.strip() else "No readable cells were found in the XLSX file."
        except Exception as exc:  # noqa: BLE001
            return "", f"XLSX could not be parsed: {exc}"

    @staticmethod
    def _decode_text_payload(raw_bytes: bytes) -> str:
        if not raw_bytes:
            return ""
        if b"\x00" in raw_bytes:
            return ""

        for encoding in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        return raw_bytes.decode("utf-8", errors="ignore")

    @staticmethod
    def _sanitize_file_name(name: str) -> str:
        base_name = Path(name).name.strip()
        if not base_name:
            return "upload.txt"

        collapsed = re.sub(r"\s+", "-", base_name)
        safe = re.sub(r"[^A-Za-z0-9._-]", "_", collapsed)
        if "." not in safe:
            safe = f"{safe}.txt"
        return safe[:120]
